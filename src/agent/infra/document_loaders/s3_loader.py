# src/stock_assistant/infra/document_loaders/s3_loader.py
import boto3
import io
import asyncio
from typing import List, Optional, Dict, Any
from pathlib import Path
import PyPDF2
from docx import Document as DocxDocument
from botocore.exceptions import ClientError
from .base import BaseDocumentLoader
from ..base import HealthCheckable
from ...domain.entities.document import DocumentChunk, DocumentMetadata
from ...shared.settings.settings import settings
from ...shared.exceptions.domain_exceptions import DocumentProcessingError
from ...shared.logging.logger import Logger
logger = Logger.get_logger(__name__)


class S3DocumentLoader(BaseDocumentLoader, HealthCheckable):
    """Load and process documents from S3"""
    
    def __init__(self):
        self.session = boto3.Session(
            aws_access_key_id=settings.s3.aws_access_key_id,
            aws_secret_access_key=settings.s3.aws_secret_access_key,
            region_name=settings.s3.aws_region
        )
        self.s3_client = self.session.client('s3')
        self.bucket_name = settings.s3.bucket_name
        self.documents_prefix = settings.s3.documents_prefix
        self.supported_extensions = {'.pdf', '.docx', '.txt', '.doc'}


    async def check_health(self) -> Dict[str, Any]:
        """Check the health of S3 document loader."""
        try:
            # Test bucket accessibility with head_bucket
            def _check_bucket():
                self.s3_client.head_bucket(Bucket=self.bucket_name)
                return True
            
            await asyncio.get_event_loop().run_in_executor(None, _check_bucket)
            
            return {
                "status": "healthy",
                "details": {
                    "message": f"S3 bucket {self.bucket_name} is accessible",
                    "bucket_name": self.bucket_name,
                    "region": settings.s3.aws_region
                }
            }
        except ClientError as e:
            error_code = e.response["Error"]["Code"]
            error_message = e.response["Error"]["Message"]
            logger.error(f"S3 health check failed: {error_code} - {error_message}")
            return {
                "status": "unhealthy",
                "details": {
                    "error": f"{error_code}: {error_message}",
                    "message": f"Failed to access S3 bucket {self.bucket_name}"
                }
            }
        except Exception as e:
            logger.error(f"S3 health check failed: {str(e)}")
            return {
                "status": "unhealthy",
                "details": {
                    "error": str(e),
                    "message": f"Failed to check S3 bucket {self.bucket_name}"
                }
            }

    
    async def list_documents(self, prefix: Optional[str] = None) -> List[Dict[str, Any]]:
        """List all documents in S3 bucket with metadata"""
        try:
            search_prefix = prefix or self.documents_prefix
            
            # Use asyncio to run the S3 operation in a thread
            def _list_objects():
                paginator = self.s3_client.get_paginator('list_objects_v2')
                page_iterator = paginator.paginate(
                    Bucket=self.bucket_name,
                    Prefix=search_prefix
                )
                
                documents = []
                for page in page_iterator:
                    if 'Contents' in page:
                        for obj in page['Contents']:
                            key = obj['Key']
                            if self._is_supported_document(key):
                                documents.append({
                                    'key': key,
                                    'size': obj['Size'],
                                    'last_modified': obj['LastModified'],
                                    'etag': obj['ETag'].strip('"')
                                })
                return documents
            
            documents = await asyncio.get_event_loop().run_in_executor(None, _list_objects)
            
            logger.info(f"Found {len(documents)} supported documents in S3")
            return documents
            
        except Exception as e:
            logger.error(f"Failed to list S3 documents: {str(e)}")
            raise DocumentProcessingError(f"Failed to list documents: {str(e)}")
    
    async def load_and_chunk_document(self, s3_key: str) -> List[DocumentChunk]:
        """Load document from S3 and split into chunks"""
        try:
            # Download document from S3
            def _download_file():
                response = self.s3_client.get_object(Bucket=self.bucket_name, Key=s3_key)
                return response['Body'].read(), response.get('LastModified')
            
            file_content, last_modified = await asyncio.get_event_loop().run_in_executor(None, _download_file)
            
            # Determine document type and extract text
            file_extension = Path(s3_key).suffix.lower()
            document_text = await self._extract_text(file_content, file_extension, s3_key)
            
            if not document_text or len(document_text.strip()) < 50:
                logger.warning(f"Document {s3_key} has insufficient content, skipping")
                return []
            
            # Create metadata
            metadata = DocumentMetadata(
                source=f"s3://{self.bucket_name}/{s3_key}",
                title=self._extract_title_from_path(s3_key),
                created_date=last_modified,
                document_type=file_extension[1:] if file_extension else "unknown",
                tags=self._extract_tags_from_path(s3_key)
            )
            
            # Split into chunks
            chunks = await self._split_text_into_chunks(document_text, metadata)
            
            logger.info(f"Processed document {s3_key} into {len(chunks)} chunks")
            return chunks
            
        except Exception as e:
            logger.error(f"Failed to process document {s3_key}: {str(e)}")
            raise DocumentProcessingError(f"Failed to process document {s3_key}: {str(e)}")
    
    async def load_all_documents(self) -> List[DocumentChunk]:
        """Load and process all documents from S3"""
        try:
            documents = await self.list_documents()
            all_chunks = []
            
            # Process documents in batches to avoid overwhelming the system
            batch_size = 5
            for i in range(0, len(documents), batch_size):
                batch = documents[i:i + batch_size]
                
                # Process batch concurrently
                batch_tasks = []
                for doc in batch:
                    task = self.load_and_chunk_document(doc['key'])
                    batch_tasks.append(task)
                
                try:
                    batch_results = await asyncio.gather(*batch_tasks, return_exceptions=True)
                    
                    for doc, result in zip(batch, batch_results):
                        if isinstance(result, Exception):
                            logger.warning(f"Failed to process {doc['key']}: {str(result)}")
                        else:
                            all_chunks.extend(result)
                            
                except Exception as e:
                    logger.error(f"Batch processing failed: {str(e)}")
                    continue
                
                # Small delay between batches
                await asyncio.sleep(0.5)
            
            logger.info(f"Loaded {len(all_chunks)} total chunks from {len(documents)} documents")
            return all_chunks
            
        except Exception as e:
            logger.error(f"Failed to load all documents: {str(e)}")
            raise DocumentProcessingError(f"Failed to load documents: {str(e)}")
    
    async def _extract_text(self, file_content: bytes, file_extension: str, s3_key: str) -> str:
        """Extract text content based on file type"""
        try:
            if file_extension == '.pdf':
                return await self._extract_pdf_text(file_content)
            elif file_extension == '.docx':
                return await self._extract_docx_text(file_content)
            elif file_extension == '.doc':
                return await self._extract_doc_text(file_content)
            elif file_extension == '.txt':
                return await self._extract_txt_text(file_content)
            else:
                raise DocumentProcessingError(f"Unsupported file type: {file_extension}")
        except Exception as e:
            logger.error(f"Text extraction failed for {s3_key}: {str(e)}")
            raise DocumentProcessingError(f"Text extraction failed: {str(e)}")
    
    async def _extract_pdf_text(self, file_content: bytes) -> str:
        """Extract text from PDF"""
        def _extract():
            try:
                pdf_file = io.BytesIO(file_content)
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                
                text = ""
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                    except Exception as e:
                        logger.warning(f"Failed to extract text from PDF page {page_num + 1}: {str(e)}")
                        continue
                
                return text.strip()
            except Exception as e:
                raise DocumentProcessingError(f"PDF extraction failed: {str(e)}")
        
        return await asyncio.get_event_loop().run_in_executor(None, _extract)
    
    async def _extract_docx_text(self, file_content: bytes) -> str:
        """Extract text from DOCX"""
        def _extract():
            try:
                doc_file = io.BytesIO(file_content)
                doc = DocxDocument(doc_file)
                
                text = ""
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text += paragraph.text + "\n"
                
                # Also extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            text += " | ".join(row_text) + "\n"
                
                return text.strip()
            except Exception as e:
                raise DocumentProcessingError(f"DOCX extraction failed: {str(e)}")
        
        return await asyncio.get_event_loop().run_in_executor(None, _extract)
    
    async def _extract_doc_text(self, file_content: bytes) -> str:
        """Extract text from DOC (legacy Word format)"""
        # For .doc files, we'll try a simple approach or return a warning
        logger.warning("DOC format not fully supported, consider converting to DOCX")
        return "Document format not fully supported. Please convert to DOCX format."
    
    async def _extract_txt_text(self, file_content: bytes) -> str:
        """Extract text from TXT"""
        def _extract():
            try:
                # Try multiple encodings
                encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
                
                for encoding in encodings:
                    try:
                        return file_content.decode(encoding)
                    except UnicodeDecodeError:
                        continue
                
                # If all encodings fail, use utf-8 with error handling
                return file_content.decode('utf-8', errors='replace')
                
            except Exception as e:
                raise DocumentProcessingError(f"TXT extraction failed: {str(e)}")
        
        return await asyncio.get_event_loop().run_in_executor(None, _extract)
    
    async def _split_text_into_chunks(
        self, 
        text: str, 
        metadata: DocumentMetadata
    ) -> List[DocumentChunk]:
        """Split text into smaller chunks for vector storage"""
        chunk_size = settings.embeddings.chunk_size
        chunk_overlap = settings.embeddings.chunk_overlap
        
        # Clean text first
        text = self._clean_text(text)
        
        if len(text) < 100:  # Skip very short texts
            return []
        
        chunks = []
        start = 0
        chunk_index = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # Try to find a good breaking point
            if end < len(text):
                # Look for sentence endings near the chunk boundary
                search_range = min(200, len(text) - end)
                best_break = end
                
                for i in range(search_range):
                    char = text[end + i]
                    if char in '.!?':
                        # Check if it's likely end of sentence (followed by space/newline)
                        if (end + i + 1 < len(text) and 
                            text[end + i + 1] in ' \n\r'):
                            best_break = end + i + 1
                            break
                    elif char in '\n':
                        best_break = end + i + 1
                        break
                
                end = best_break
            
            chunk_text = text[start:end].strip()
            
            if len(chunk_text) > 50:  # Only add meaningful chunks
                chunk = DocumentChunk(
                    content=chunk_text,
                    metadata=metadata,
                    chunk_index=chunk_index,
                    start_char=start,
                    end_char=end
                )
                chunks.append(chunk)
                chunk_index += 1
            
            # Move start position with overlap
            start = max(start + 1, end - chunk_overlap)
            
            # Prevent infinite loop
            if start >= len(text):
                break
                
        return chunks
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        # Remove excessive whitespace
        import re
        text = re.sub(r'\n\s*\n', '\n\n', text)  # Multiple newlines to double
        text = re.sub(r' +', ' ', text)  # Multiple spaces to single
        text = re.sub(r'\t+', ' ', text)  # Tabs to spaces
        
        # Remove very long lines of repeated characters (often formatting artifacts)
        text = re.sub(r'(.)\1{20,}', '', text)
        
        return text.strip()
    
    def _extract_title_from_path(self, s3_key: str) -> str:
        """Extract a readable title from S3 key"""
        filename = Path(s3_key).stem
        # Replace underscores and hyphens with spaces
        title = filename.replace('_', ' ').replace('-', ' ')
        # Capitalize words
        title = ' '.join(word.capitalize() for word in title.split())
        return title
    
    def _extract_tags_from_path(self, s3_key: str) -> List[str]:
        """Extract tags from S3 path structure"""
        tags = []
        path_parts = s3_key.split('/')
        
        # Add folder names as tags
        for part in path_parts[:-1]:  # Exclude filename
            if part and part != self.documents_prefix.rstrip('/'):
                tags.append(part.replace('_', ' ').replace('-', ' ').title())
        
        # Add file extension as tag
        ext = Path(s3_key).suffix.lower()
        if ext:
            tags.append(ext[1:].upper())
        
        return tags
    
    def _is_supported_document(self, key: str) -> bool:
        """Check if document type is supported"""
        if key.endswith('/'):  # Skip directories
            return False
            
        extension = Path(key).suffix.lower()
        return extension in self.supported_extensions
    
    async def get_document_info(self, s3_key: str) -> Dict[str, Any]:
        """Get information about a specific document"""
        try:
            def _get_info():
                response = self.s3_client.head_object(Bucket=self.bucket_name, Key=s3_key)
                return {
                    'key': s3_key,
                    'size': response['ContentLength'],
                    'last_modified': response['LastModified'],
                    'content_type': response.get('ContentType', 'unknown'),
                    'etag': response['ETag'].strip('"')
                }
            
            return await asyncio.get_event_loop().run_in_executor(None, _get_info)
            
        except Exception as e:
            logger.error(f"Failed to get document info for {s3_key}: {str(e)}")
            raise DocumentProcessingError(f"Failed to get document info: {str(e)}")
    
    async def check_document_exists(self, s3_key: str) -> bool:
        """Check if document exists in S3"""
        try:
            def _check():
                try:
                    self.s3_client.head_object(Bucket=self.bucket_name, Key=s3_key)
                    return True
                except self.s3_client.exceptions.NoSuchKey:
                    return False
            
            return await asyncio.get_event_loop().run_in_executor(None, _check)
            
        except Exception as e:
            logger.error(f"Error checking document existence {s3_key}: {str(e)}")
            return False